#!/usr/bin/env python3
"""Convert MOU HTML to DOCX preserving formatting."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString, Tag
import copy

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0xBA, 0x94, 0x52)
INK = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x4A, 0x51, 0x58)
GRAY = RGBColor(0x8A, 0x8F, 0x96)
LIGHT_BG = RGBColor(0xED, 0xEA, 0xE2)


def set_font(run, name='EB Garamond', size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = color


def para_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def add_horizontal_rule(doc, color=RGBColor(0xE3, 0xE0, 0xD8)):
    """Add a bottom border to the last paragraph as a rule."""
    p = doc.add_paragraph()
    para_spacing(p, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E3E0D8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=24, after=8)
    # Bottom border on heading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E3E0D8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, 'EB Garamond', size=15, bold=False, color=NAVY)
    return p


def get_text_with_formatting(element, para):
    """Recursively add text from element into paragraph with formatting."""
    if isinstance(element, NavigableString):
        text = str(element)
        # Decode HTML entities
        text = text.replace('’', "'").replace('‘', "'")
        text = text.replace('“', '"').replace('”', '"')
        text = text.replace('–', '–').replace('—', '—')
        text = text.replace(' ', ' ')
        text = text.replace('®', '®').replace('©', '©')
        if text:
            run = para.add_run(text)
            set_font(run, 'EB Garamond', size=11, color=INK)
        return

    tag = element.name
    if tag is None:
        return

    if tag in ('strong', 'b'):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).replace(' ', ' ')
                if text:
                    run = para.add_run(text)
                    set_font(run, 'EB Garamond', size=11, bold=True, color=INK)
            else:
                # Check if child is span with highlight
                inner_style = child.get('style', '')
                if 'background:#EDEAE2' in inner_style or 'background: #EDEAE2' in inner_style:
                    inner_text = child.get_text()
                    if inner_text:
                        run = para.add_run(inner_text)
                        set_font(run, 'EB Garamond', size=11, bold=True, color=INK)
                        # Yellow highlight to indicate placeholder
                        rPr = run._r.get_or_add_rPr()
                        highlight = OxmlElement('w:highlight')
                        highlight.set(qn('w:val'), 'yellow')
                        rPr.append(highlight)
                else:
                    get_text_with_formatting(child, para)
    elif tag in ('em', 'i'):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).replace(' ', ' ')
                if text:
                    run = para.add_run(text)
                    set_font(run, 'EB Garamond', size=11, italic=True, color=INK)
            else:
                get_text_with_formatting(child, para)
    elif tag == 'span':
        style = element.get('style', '')
        is_placeholder = 'background:#EDEAE2' in style or 'background: #EDEAE2' in style
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).replace(' ', ' ')
                if text:
                    run = para.add_run(text)
                    set_font(run, 'EB Garamond', size=11, color=INK)
                    if is_placeholder:
                        rPr = run._r.get_or_add_rPr()
                        highlight = OxmlElement('w:highlight')
                        highlight.set(qn('w:val'), 'yellow')
                        rPr.append(highlight)
            else:
                get_text_with_formatting(child, para)
    elif tag == 'br':
        run = para.add_run('\n')
    else:
        for child in element.children:
            get_text_with_formatting(child, para)


def process_paragraph(doc, elem, indent=False):
    """Process a <p> tag into a doc paragraph."""
    p = doc.add_paragraph()
    style = elem.get('style', '')
    if 'text-align:center' in style or 'text-align: center' in style:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before=0, after=6)
    p.paragraph_format.line_spacing = Pt(17)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    get_text_with_formatting(elem, p)
    return p


def process_list_item(doc, li_elem, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    para_spacing(p, before=0, after=4)
    p.paragraph_format.line_spacing = Pt(17)
    get_text_with_formatting(li_elem, p)
    return p


def add_header_footer(doc):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.clear()
    # Left text
    r1 = hp.add_run("SUAPS · Yale “Thinking about the Impossible” MOU — Draft v0.1")
    r1.font.name = 'Source Sans Pro'
    r1.font.size = Pt(7)
    r1.font.color.rgb = GRAY
    # Tab to right
    tab_stop_xml = '\t'
    r_tab = hp.add_run('\t')
    r_tab.font.size = Pt(7)
    # Right text
    r2 = hp.add_run("Confidential — Not for External Circulation")
    r2.font.name = 'Source Sans Pro'
    r2.font.size = Pt(7)
    r2.font.color.rgb = GRAY
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Set tab stop to right-align the right text
    pPr = hp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    # Page width minus margins = 8.5 - 1.7 = 6.8 inches = 6120 twips
    tab.set(qn('w:pos'), '6120')
    tabs.append(tab)
    pPr.append(tabs)
    # Bottom border on header paragraph
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E3E0D8')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    r3 = fp.add_run("Society for UAP Studies")
    r3.font.name = 'Source Sans Pro'
    r3.font.size = Pt(7)
    r3.font.color.rgb = GRAY
    r_tab2 = fp.add_run('\t')
    r_tab2.font.size = Pt(7)
    r4 = fp.add_run("societyforuapstudies.org")
    r4.font.name = 'Source Sans Pro'
    r4.font.size = Pt(7)
    r4.font.color.rgb = GRAY
    # Tab stop
    pPr2 = fp._p.get_or_add_pPr()
    tabs2 = OxmlElement('w:tabs')
    tab2 = OxmlElement('w:tab')
    tab2.set(qn('w:val'), 'right')
    tab2.set(qn('w:pos'), '6120')
    tabs2.append(tab2)
    pPr2.append(tabs2)
    # Top border on footer
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top')
    top2.set(qn('w:val'), 'single')
    top2.set(qn('w:sz'), '4')
    top2.set(qn('w:space'), '1')
    top2.set(qn('w:color'), 'E3E0D8')
    pBdr2.append(top2)
    pPr2.append(pBdr2)

    # Page number in center of footer
    # Actually let's do it simpler - just left/right like original
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in kwargs.items():
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_cover_block(doc):
    """Add the Internal Draft status block as a shaded table."""
    # Title label
    label_p = doc.add_paragraph()
    para_spacing(label_p, before=0, after=4)
    run = label_p.add_run("INTERNAL DRAFT — GOVERNANCE & CIRCULATION NOTE")
    run.font.name = 'Source Sans Pro'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = GOLD

    # Table with 2 columns: label | content
    rows_data = [
        ("STATUS", "v0.1 — First draft for internal legal/programmatic review. Not yet circulated to counterparty."),
        ("PREPARED BY", "Office of the Executive Director, Society for UAP Studies"),
        ("AUTHORIZATION BASIS",
         "Executed under the Executive Director’s ordinary executive authority to commit programmatic and "
         "sponsorship funds within Board-delegated financial thresholds (Board of Directors Governance Manual §7, "
         "Board–CEO Relationship), consistent with prior sponsorship instruments signed by the Executive Director "
         "alone (e.g., Argentina Congress, $2,000; Brazil Event, $700). As SUAPS’s first agreement of "
         "‘partnership’ character rather than simple event sponsorship, this Agreement will be noted to the "
         "Board at its next regularly scheduled meeting rather than requiring advance Board approval, consistent with "
         "the $5,000 commitment falling well under the $100,000 threshold requiring Board action under the DDRG "
         "delegation framework."),
        ("OPEN ITEM",
         "Yale-side counterparty entity is unconfirmed. This draft places a bracketed placeholder wherever the exact "
         "legal name, entity type, and authorized signatory must be inserted. Recommend confirming whether Sri Tata "
         "is signing on behalf of a recognized Yale student organization, a Yale department/fiscal sponsor, or in an "
         "individual capacity, before this MOU is circulated externally — the answer changes both the signature "
         "block and Yale’s own trademark/licensing sign-off requirements (see Section 7)."),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    # Remove outer borders
    for i, (label, content) in enumerate(rows_data):
        row = table.rows[i]
        # Label cell
        lc = row.cells[0]
        shade_cell(lc, 'EDEAE2')
        lc.width = Inches(1.1)
        lp = lc.paragraphs[0]
        lp.clear()
        lr = lp.add_run(label)
        lr.font.name = 'Source Sans Pro'
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.color.rgb = GRAY
        lp.paragraph_format.space_after = Pt(0)

        # Content cell
        cc = row.cells[1]
        shade_cell(cc, 'EDEAE2')
        cp = cc.paragraphs[0]
        cp.clear()
        if label == "OPEN ITEM":
            # "Yale-side counterparty entity is unconfirmed." in bold
            bold_part = "Yale-side counterparty entity is unconfirmed."
            cr_bold = cp.add_run(bold_part)
            cr_bold.font.name = 'Source Sans Pro'
            cr_bold.font.size = Pt(9)
            cr_bold.font.bold = True
            cr_bold.font.color.rgb = INK
            rest = content[len(bold_part):]
            cr_rest = cp.add_run(rest)
            cr_rest.font.name = 'Source Sans Pro'
            cr_rest.font.size = Pt(9)
            cr_rest.font.color.rgb = INK
        else:
            cr = cp.add_run(content)
            cr.font.name = 'Source Sans Pro'
            cr.font.size = Pt(9)
            cr.font.color.rgb = INK
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = Pt(13)

    # Add a left gold border via paragraph border workaround - use a framing paragraph
    # Actually we'll add a colored left-border paragraph before the table
    spacer = doc.add_paragraph()
    para_spacing(spacer, before=8, after=4)

    return table


def add_signature_table(doc):
    """Add the two-column signature table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    # Remove all borders
    def no_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    row = table.rows[0]
    left_cell = row.cells[0]
    spacer_cell = row.cells[1]
    right_cell = row.cells[2]

    no_borders(left_cell)
    no_borders(spacer_cell)
    no_borders(right_cell)

    left_cell.width = Inches(3.0)
    spacer_cell.width = Inches(0.3)
    right_cell.width = Inches(3.0)

    def fill_sig_cell(cell, party_line, name_line, title_line):
        # Signature line
        sig_p = cell.paragraphs[0]
        sig_p.clear()
        pPr = sig_p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '333333')
        pBdr.append(bottom)
        pPr.append(pBdr)
        sig_p.paragraph_format.space_before = Pt(24)
        sig_p.paragraph_format.space_after = Pt(4)

        # Party line
        party_p = cell.add_paragraph()
        party_p.paragraph_format.space_before = Pt(2)
        party_p.paragraph_format.space_after = Pt(2)
        for part in party_line:
            if isinstance(part, tuple):
                text, is_placeholder = part
                r = party_p.add_run(text)
                r.font.name = 'EB Garamond'
                r.font.size = Pt(11)
                if is_placeholder:
                    rPr = r._r.get_or_add_rPr()
                    highlight = OxmlElement('w:highlight')
                    highlight.set(qn('w:val'), 'yellow')
                    rPr.append(highlight)
            else:
                r = party_p.add_run(part)
                r.font.name = 'EB Garamond'
                r.font.size = Pt(11)

        # Name line
        name_p = cell.add_paragraph()
        name_p.paragraph_format.space_before = Pt(1)
        name_p.paragraph_format.space_after = Pt(2)
        for part in name_line:
            if isinstance(part, tuple):
                text, is_placeholder = part
                r = name_p.add_run(text)
                r.font.name = 'EB Garamond'
                r.font.size = Pt(11)
                r.font.color.rgb = MUTED
                if is_placeholder:
                    rPr = r._r.get_or_add_rPr()
                    highlight = OxmlElement('w:highlight')
                    highlight.set(qn('w:val'), 'yellow')
                    rPr.append(highlight)
            else:
                r = name_p.add_run(part)
                r.font.name = 'EB Garamond'
                r.font.size = Pt(11)
                r.font.color.rgb = MUTED

        # Date line
        date_p = cell.add_paragraph()
        date_p.paragraph_format.space_before = Pt(10)
        date_p.paragraph_format.space_after = Pt(2)
        dr = date_p.add_run(title_line)
        dr.font.name = 'EB Garamond'
        dr.font.size = Pt(11)
        dr.font.color.rgb = MUTED

    fill_sig_cell(
        left_cell,
        ["For the Society for UAP Studies"],
        ["Dr. Michael C. Cifone\nFounding Executive Director & President"],
        "Date: ______________________"
    )
    fill_sig_cell(
        right_cell,
        ["For ", ("[Organizer — legal name TBD]", True)],
        ["Sri Tata\n", ("[Title/role TBD]", True)],
        "Date: ______________________"
    )

    spacer = doc.add_paragraph()
    para_spacing(spacer, before=0, after=0)


def build_doc():
    doc = Document()

    # Page setup: letter, margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'EB Garamond'
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # Header/footer
    add_header_footer(doc)

    # ===== COVER BLOCK =====
    add_cover_block(doc)

    # ===== TITLE =====
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(org_p, before=16, after=4)
    org_r = org_p.add_run("The Society for UAP Studies")
    org_r.font.name = 'Source Sans Pro'
    org_r.font.size = Pt(9)
    org_r.font.bold = True
    org_r.font.color.rgb = GOLD

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(title_p, before=4, after=4)
    title_r = title_p.add_run("Memorandum of Understanding")
    title_r.font.name = 'EB Garamond'
    title_r.font.size = Pt(24)
    title_r.font.bold = False
    title_r.font.color.rgb = NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(subtitle_p, before=0, after=12)
    subtitle_r = subtitle_p.add_run(
        'Academic Partnership & Event Sponsorship — “Thinking about the Impossible”'
    )
    subtitle_r.font.name = 'EB Garamond'
    subtitle_r.font.size = Pt(12)
    subtitle_r.font.italic = True
    subtitle_r.font.color.rgb = MUTED

    between_p = doc.add_paragraph()
    between_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(between_p, before=4, after=4)
    between_r = between_p.add_run("— between —")
    between_r.font.name = 'EB Garamond'
    between_r.font.size = Pt(11)
    between_r.font.color.rgb = MUTED

    # Parties
    parties_p = doc.add_paragraph()
    parties_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(parties_p, before=4, after=16)
    parties_p.paragraph_format.line_spacing = Pt(18)

    def add_run_parties(text, bold=False, italic=False, placeholder=False, color=None):
        r = parties_p.add_run(text)
        r.font.name = 'EB Garamond'
        r.font.size = Pt(11)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
        if placeholder:
            rPr = r._r.get_or_add_rPr()
            h = OxmlElement('w:highlight')
            h.set(qn('w:val'), 'yellow')
            rPr.append(h)

    add_run_parties("The Society for UAP Studies", bold=True)
    add_run_parties(' (“')
    add_run_parties("SUAPS", bold=True)
    add_run_parties('” or the “')
    add_run_parties("Society", bold=True)
    add_run_parties(
        '”), a non-profit organization registered in California and incorporated in the State of Delaware, United States of America\n\nand\n\n'
    )
    add_run_parties("[YALE-SIDE COUNTERPARTY — LEGAL NAME TO BE CONFIRMED]", bold=True, placeholder=True)
    add_run_parties(
        ', acting through Mr. Sri Tata, on behalf of the organizing committee of the conference “Thinking about the Impossible” at Yale University (hereafter, the “'
    )
    add_run_parties("Organizer", bold=True)
    add_run_parties('”)')

    # Main text paragraphs
    paras = [
        ('p_body',
         'Together, SUAPS and the Organizer are referred to as the “',
         [('Parties', True), ('” and individually as a “', False), ('Party', True),
          ('.” The Parties acknowledge that this Memorandum of Understanding (the “', False),
          ('Agreement', True),
          ('”) establishes both (a) a limited financial sponsorship of the Event defined below, and '
           '(b) a scholarly collaboration between SUAPS-affiliated researchers and the Organizer’s '
           'academic program, and that these two elements are governed by distinct terms as set out in '
           'Sections 4 and 5. This Agreement does not constitute a research grant, the administration of '
           'research funds, or a joint venture, partnership, or agency relationship in the legal sense '
           'between the Parties.', False)]
         ),
    ]

    def body_para(text_parts=None, simple_text=None, align=None):
        p = doc.add_paragraph()
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=0, after=6)
        p.paragraph_format.line_spacing = Pt(17)
        if simple_text:
            r = p.add_run(simple_text)
            set_font(r, 'EB Garamond', size=11, color=INK)
        elif text_parts:
            for part in text_parts:
                if isinstance(part, tuple) and len(part) == 2:
                    text, is_bold = part
                    r = p.add_run(text)
                    set_font(r, 'EB Garamond', size=11, bold=is_bold, color=INK)
                else:
                    r = p.add_run(part)
                    set_font(r, 'EB Garamond', size=11, color=INK)
        return p

    def bold(text):
        return (text, True)

    def norm(text):
        return (text, False)

    # Parties text
    body_para(text_parts=[
        norm('Together, SUAPS and the Organizer are referred to as the “'),
        bold('Parties'), norm('” and individually as a “'), bold('Party'),
        norm('.” The Parties acknowledge that this Memorandum of Understanding (the “'),
        bold('Agreement'),
        norm('”) establishes both (a) a limited financial sponsorship of the Event defined below, and '
             '(b) a scholarly collaboration between SUAPS-affiliated researchers and the Organizer’s '
             'academic program, and that these two elements are governed by distinct terms as set out in '
             'Sections 4 and 5. This Agreement does not constitute a research grant, the administration of '
             'research funds, or a joint venture, partnership, or agency relationship in the legal sense '
             'between the Parties.')
    ])

    body_para(simple_text='Each Party represents that the individual signing this Agreement on its behalf has been duly authorized to do so.')

    # === SECTION 1 ===
    add_heading2(doc, '1. Purpose and Background')

    body_para(text_parts=[
        bold('1.1'), norm(' SUAPS is a non-profit academic organization committed to advancing high-quality research, '
                          'education, and dialogue regarding unidentified anomalous phenomena (UAP) and related topics, '
                          'including through its Strategic Institutional Research Initiatives (SIRIs) such as the '
                          'Empirical Weird Initiative (“'), bold('EWI'), norm('”).'),
    ])

    body_para(text_parts=[
        bold('1.2'), norm(' The Organizer is convening “Thinking about the Impossible” (the “'),
        bold('Event'), norm('”), a two-day academic conference to be held December 4–5, 2026 at Yale University, '
                            'New Haven, Connecticut, featuring approximately 8–10 speakers per day across two days, '
                            'including morning and evening keynote sessions and two panel discussions per day.'),
    ])

    body_para(text_parts=[
        bold('1.3'), norm(' The Parties wish to formalize (a) SUAPS’s financial support for the Event, and '
                          '(b) a scholarly collaboration through which SUAPS-affiliated researchers may contribute a '
                          'panel discussion drawn from the EWI’s ongoing research program, on terms that keep these '
                          'two forms of support clearly and contractually distinct.'),
    ])

    # === SECTION 2 ===
    add_heading2(doc, '2. Event Description')

    body_para(text_parts=[bold('2.1 Event Title.'), norm(' “Thinking about the Impossible”.')])
    body_para(text_parts=[bold('2.2 Dates.'), norm(' December 4–5, 2026.')])
    body_para(text_parts=[bold('2.3 Location.'), norm(' Yale University, New Haven, Connecticut (venue to be confirmed by the Organizer).')])
    body_para(text_parts=[
        bold('2.4 Format.'), norm(' In-person conference with a pre-event virtual lecture series for remote speakers; '
                                  'the Parties do not presently intend a live-streamed main event, consistent with the '
                                  'Organizer’s preference for a professional, non-interactive virtual component '
                                  '(e.g., recorded or single-direction streaming) over open interactive platforms.'),
    ])
    body_para(text_parts=[
        bold('2.5 Independent Collaborations.'), norm(' The Parties acknowledge that the Organizer maintains separate '
                                                      'relationships with other supporters of the Event, including the Soul Foundation, '
                                                      'which are managed independently by the Organizer and fall outside the scope of this Agreement.'),
    ])

    # === SECTION 3 ===
    add_heading2(doc, '3. Nature of the Relationship')

    body_para(text_parts=[
        bold('3.1'), norm(' This Agreement reflects two independent and severable forms of collaboration: a fixed '
                          'financial contribution to the Event (Section 4) and a scholarly programming collaboration '
                          'centered on the EWI (Section 5). Neither is offered, conditioned upon, or to be construed '
                          'as consideration for the other.'),
    ])

    body_para(text_parts=[
        bold('3.2'), norm('  '), bold('No Pay-for-Programming Representation.'),
        norm(' For the avoidance of doubt, and because both Parties wish to avoid any appearance that programming access '
             'has been purchased: SUAPS’s financial contribution under Section 4 confers no right, expectation, '
             'or entitlement to program placement, speaking slots, or any other academic or programmatic benefit at the '
             'Event. Any participation of SUAPS-affiliated researchers in the Event program is granted solely on the '
             'Organizer’s independent assessment of scholarly merit and fit, under the Organizer’s ordinary '
             'process for selecting speakers and panelists, as further described in Section 5.'),
    ])

    # === SECTION 4 ===
    add_heading2(doc, '4. Financial Contribution')

    body_para(text_parts=[
        bold('4.1 Amount.'), norm(' SUAPS agrees to contribute a total of USD 5,000 (five thousand United States dollars) (the “'),
        bold('Contribution'), norm('”) in support of the Event.'),
    ])
    body_para(text_parts=[
        bold('4.2 Payment Schedule.'), norm(' The Contribution shall be paid in two installments against invoices issued by the Organizer:'),
    ])

    # List items for 4.2
    li1 = doc.add_paragraph(style='List Bullet')
    para_spacing(li1, before=0, after=4)
    li1.paragraph_format.line_spacing = Pt(17)
    li1.paragraph_format.left_indent = Inches(0.4)
    li1.paragraph_format.first_line_indent = Inches(-0.2)
    r = li1.add_run('USD 2,500 within thirty (30) days of the execution of this Agreement; and')
    set_font(r, 'EB Garamond', size=11, color=INK)

    li2 = doc.add_paragraph(style='List Bullet')
    para_spacing(li2, before=0, after=4)
    li2.paragraph_format.line_spacing = Pt(17)
    li2.paragraph_format.left_indent = Inches(0.4)
    li2.paragraph_format.first_line_indent = Inches(-0.2)
    r = li2.add_run('USD 2,500 no later than thirty (30) days prior to the start of the Event (i.e., on or before November 4, 2026).')
    set_font(r, 'EB Garamond', size=11, color=INK)

    body_para(simple_text='Payment timing is tied solely to the calendar milestones above and is not conditioned on, or linked to, any programmatic decision, including the panel described in Section 5.')

    body_para(text_parts=[
        bold('4.3 Invoicing.'), norm(' Each invoice shall include: the legal name and address of the Organizer (or its fiscal sponsor); '
                                     'a clear description (e.g., “Sponsorship of ‘Thinking about the Impossible,’ December 4–5, 2026”); '
                                     'the amount, currency, and due date; and complete banking or payment details. SUAPS shall not be in default '
                                     'of any payment obligation until thirty (30) days have elapsed following its receipt of both a valid invoice '
                                     'and complete written banking instructions.'),
    ])
    body_para(text_parts=[
        bold('4.4 Recognition.'), norm(' In consideration of the Contribution, the Organizer will recognize SUAPS as a named academic partner '
                                       'and supporter of the Event in event materials (website, program, signage), at a level consistent with the '
                                       'Contribution amount and comparable to recognition afforded to organizations of similar support level. The '
                                       'specific form of recognition shall be confirmed in writing between the Parties no later than thirty (30) days before the Event.'),
    ])
    body_para(text_parts=[
        bold('4.5 Records.'), norm(' The Organizer shall provide a receipt and/or written confirmation of each payment. Each Party will retain '
                                   'invoices, receipts, and supporting documentation for internal audit, compliance, and donor-reporting purposes.'),
    ])

    # === SECTION 5 ===
    add_heading2(doc, '5. Academic Collaboration — Empirical Weird Initiative Panel')

    body_para(text_parts=[
        bold('5.1 Proposed Contribution.'), norm(' SUAPS proposes to contribute one (1) panel discussion at the Event drawn from its '
                                                  'Empirical Weird Initiative (EWI), a Strategic Institutional Research Initiative examining '
                                                  'anomalous experience through an interdisciplinary, empirically grounded methodology.'),
    ])
    body_para(text_parts=[
        bold('5.2 Illustrative Roster.'), norm(' The following SUAPS-affiliated researchers are put forward as a resource pool for the '
                                                'Organizer’s consideration. This roster is illustrative and non-binding; final panel '
                                                'composition, format, title, and scheduling remain within the Organizer’s sole discretion, '
                                                'exercised in consultation with the EWI Project Lead:'),
    ])

    roster = [
        ('Dr. Kimberly Engels', ' — Molloy University; John Mack Institute; SUAPS Advisory Board ', '(EWI Project Lead)'),
        ('Dr. Michael Silberstein', ' — Elizabethtown College; SUAPS Co-Founder ', '(EWI Co-PI)'),
        ('Dr. Wesley Watters', ' — Wellesley College; SUAPS Advisory Board', None),
        ('Dr. Greg Eghigian', ' — Pennsylvania State University; SUAPS Advisory Council', None),
        ('Dr. William Seager', ' — University of Toronto (Emeritus)', None),
        ('Dr. habil. Bernd-Christian Otto', ' — Friedrich-Alexander University', None),
        ('Dr. Knut Graw', ' — Friedrich-Alexander University', None),
        ('Dr. habil. Harald Atmanspacher', ' — ETH Zurich (Emeritus)', None),
        ('Dr. Michael C. Cifone', ' — St. John’s University; SUAPS Founding Executive Director', None),
        ('Dr. Kelly Hayes', ' — Indiana University', None),
        ('Annalisa Ventola', ' — Independent Scholar; Executive Director, Parapsychological Association', None),
    ]
    for name, affil, role in roster:
        li = doc.add_paragraph(style='List Bullet')
        para_spacing(li, before=0, after=2)
        li.paragraph_format.line_spacing = Pt(16)
        li.paragraph_format.left_indent = Inches(0.4)
        li.paragraph_format.first_line_indent = Inches(-0.2)
        r1 = li.add_run(name)
        set_font(r1, 'EB Garamond', size=11, bold=True, color=INK)
        r2 = li.add_run(affil)
        set_font(r2, 'EB Garamond', size=11, color=INK)
        if role:
            r3 = li.add_run(role)
            set_font(r3, 'EB Garamond', size=11, italic=True, color=INK)

    body_para(text_parts=[
        bold('5.3 Editorial and Programmatic Independence.'), norm(' The Organizer retains full and independent authority over program content, '
                                                                     'panel composition, framing, and scheduling, applying the same standards of academic merit and fit used for all '
                                                                     'other Event programming. SUAPS will not represent to any third party that its Contribution entitles it to any '
                                                                     'specific programming outcome.'),
    ])
    body_para(text_parts=[
        bold('5.4 Confirmation.'), norm(' The Organizer will confirm the final panel roster, title, and format in writing no later than sixty '
                                        '(60) days before the Event. If the Organizer determines that a panel drawn from the EWI is not a fit for '
                                        'the final program, this Agreement’s financial terms under Section 4 remain unaffected and unconditional.'),
    ])
    body_para(text_parts=[
        bold('5.5 Speaker Views Are Their Own.'), norm(' Views expressed by SUAPS-affiliated participants at the Event are their own and do not '
                                                        'necessarily represent the official positions of SUAPS or the Organizer.'),
    ])

    # === SECTION 6 ===
    add_heading2(doc, '6. Roles and Responsibilities')

    body_para(text_parts=[bold('6.1 The Organizer will:')])

    org_items = [
        'Plan, organize, and deliver the Event;',
        'Provide the recognition described in Section 4.4;',
        'Communicate promptly any significant changes to the Event (dates, venue, format, program);',
        'Coordinate logistics, technical requirements, and scheduling for any confirmed SUAPS-affiliated speakers; and',
        'Provide SUAPS with a short post-Event report consistent with Section 8.',
    ]
    for item in org_items:
        li = doc.add_paragraph(style='List Bullet')
        para_spacing(li, before=0, after=4)
        li.paragraph_format.line_spacing = Pt(17)
        li.paragraph_format.left_indent = Inches(0.4)
        li.paragraph_format.first_line_indent = Inches(-0.2)
        r = li.add_run(item)
        set_font(r, 'EB Garamond', size=11, color=INK)

    body_para(text_parts=[bold('6.2 SUAPS will:')])

    suaps_items = [
        'Pay the Contribution per the schedule in Section 4.2;',
        'Make the illustrative roster in Section 5.2 available to the Organizer and coordinate scheduling with any confirmed participants;',
        'Provide its logo and basic brand guidelines for use per Section 7; and',
        'Respect the Organizer’s reasonable event guidelines and academic standards.',
    ]
    for item in suaps_items:
        li = doc.add_paragraph(style='List Bullet')
        para_spacing(li, before=0, after=4)
        li.paragraph_format.line_spacing = Pt(17)
        li.paragraph_format.left_indent = Inches(0.4)
        li.paragraph_format.first_line_indent = Inches(-0.2)
        r = li.add_run(item)
        set_font(r, 'EB Garamond', size=11, color=INK)

    # === SECTION 7 ===
    add_heading2(doc, '7. Branding, Communications, and Use of Names')

    body_para(text_parts=[
        bold('7.1'), norm(' Each Party grants the other a non-exclusive, royalty-free, limited license to use its name and logo '
                          'solely to recognize and promote this collaboration, subject to any brand guidelines shared in writing.'),
    ])
    body_para(text_parts=[
        bold('7.2 Yale Trademark Policies.'), norm(' The Parties acknowledge that any use of Yale University’s name, seal, or marks — '
                                                    'as distinct from the Organizer’s own event branding — is subject to Yale University’s '
                                                    'separate trademark and licensing policies and requires the Organizer to obtain any necessary approvals '
                                                    'independently of this Agreement. Nothing in this Agreement authorizes use of Yale University’s name '
                                                    'or marks beyond what the Organizer is itself authorized to grant.'),
    ])
    body_para(text_parts=[
        bold('7.3'), norm(' Any use of a Party’s name or logo outside the context of the Event, or in a manner implying endorsement '
                          'beyond the scope of this Agreement, requires that Party’s prior written approval.'),
    ])
    body_para(text_parts=[
        bold('7.4'), norm(' Public communications referencing SUAPS should refer to it by its full name (“Society for UAP Studies”) '
                          'at least once and should avoid implying that SUAPS endorses all views expressed at the Event.'),
    ])

    # === SECTION 8 ===
    add_heading2(doc, '8. Reporting')

    body_para(simple_text='Within forty-five (45) days after the conclusion of the Event, the Organizer shall provide SUAPS with a short written '
                          'report including: total attendance and a basic audience profile; confirmation of the recognition and any panel delivered '
                          'under Sections 4 and 5, noting any deviations from plan; and any notable outcomes or follow-up opportunities '
                          'relevant to SUAPS’s mission.')

    # === SECTION 9 ===
    add_heading2(doc, '9. Compliance, Ethics, and Reputational Standards')

    body_para(text_parts=[
        bold('9.1'), norm(' The Organizer shall ensure that the Event and its programming adhere to applicable laws and to reasonable '
                          'standards of academic integrity, non-discrimination, and professional conduct.'),
    ])
    body_para(text_parts=[
        bold('9.2'), norm(' If SUAPS reasonably determines that any proposed content, speaker, or associated activity would materially conflict '
                          'with its mission, ethical standards, or internal policies, it may object in writing. The Parties shall confer in good '
                          'faith to make appropriate adjustments; if no resolution is reached within a reasonable period, SUAPS may withdraw its '
                          'name, logo, and institutional association from the Event without this constituting a breach of this Agreement, though '
                          'the Contribution schedule in Section 4 shall not be affected retroactively for amounts already due.'),
    ])
    body_para(text_parts=[
        bold('9.3 Anti-Corruption.'), norm(' Each Party shall comply with applicable anti-bribery, anti-corruption, and economic sanctions '
                                           'laws in connection with this Agreement.'),
    ])

    # === SECTION 10 ===
    add_heading2(doc, '10. Term and Termination')

    body_para(text_parts=[
        bold('10.1'), norm(' This Agreement takes effect on the date of the last signature below and remains in effect until sixty (60) days '
                           'after the conclusion of the Event, unless earlier terminated under this Section or extended by mutual written agreement '
                           'of the Parties.'),
    ])
    body_para(text_parts=[
        bold('10.2'), norm(' Either Party may terminate this Agreement for cause upon written notice if the other Party engages in conduct '
                           'creating a serious and reasonable risk of material harm to the terminating Party’s reputation, legal compliance, '
                           'or ethical commitments, and such conduct is not cured (if curable) within thirty (30) days after written notice.'),
    ])
    body_para(text_parts=[
        bold('10.3'), norm(' In the event of cancellation or termination prior to the Event, the Parties will consult in good faith regarding '
                           'rescheduling, partial refund, or reallocation of the Contribution, taking into account costs already committed.'),
    ])

    # === SECTION 11 ===
    add_heading2(doc, '11. Indemnification')

    body_para(simple_text='The Organizer agrees to indemnify, defend, and hold harmless the Society for UAP Studies, its officers, directors, '
                          'employees, contractors, and representatives from and against any claims, liabilities, damages, losses, and expenses '
                          '(including reasonable legal fees) arising out of or relating to: (a) the Organizer’s contracting with any '
                          'third-party service providers; (b) event promotion or public communications conducted by the Organizer; (c) any failure '
                          'to comply with applicable laws or venue requirements; or (d) bodily injury, property damage, or other incidents occurring '
                          'in connection with the Event, except to the extent caused by SUAPS’s gross negligence or willful misconduct. Nothing '
                          'in this Agreement shall be construed as SUAPS assuming responsibility for the Organizer’s local operational, '
                          'contractual, tax, or regulatory obligations.')

    # === SECTION 12 ===
    add_heading2(doc, '12. Governing Law and Dispute Resolution')

    body_para(text_parts=[
        bold('12.1'), norm(' This Agreement shall be governed by and construed in accordance with the laws of the State of Connecticut, '
                           'United States of America, without regard to conflict-of-laws principles.'),
    ])
    body_para(text_parts=[
        bold('12.2'), norm(' Any dispute arising out of or in connection with this Agreement that cannot be resolved amicably through '
                           'good-faith discussions between designated representatives of the Parties may, by mutual agreement, be referred to '
                           'mediation or another informal dispute-resolution mechanism before either Party pursues formal legal action.'),
    ])

    # === SECTION 13 ===
    add_heading2(doc, '13. Miscellaneous')

    body_para(text_parts=[
        bold('13.1 Entire Agreement.'), norm(' This Agreement constitutes the entire understanding between the Parties regarding its subject '
                                             'matter and supersedes all prior understandings, written or oral.'),
    ])
    body_para(text_parts=[
        bold('13.2 Amendments.'), norm(' Any modification must be made in writing and signed by authorized representatives of both Parties.'),
    ])
    body_para(text_parts=[
        bold('13.3 No Partnership in the Legal Sense.'), norm(' Notwithstanding the collaborative and partnership-in-spirit character of this '
                                                               'Agreement, the Parties are independent entities, and nothing herein creates a legal partnership, joint venture, agency, '
                                                               'or employment relationship between them. Neither Party may bind the other or incur obligations on the other’s behalf '
                                                               'except as expressly set forth herein.'),
    ])
    body_para(text_parts=[
        bold('13.4 Assignment.'), norm(' Neither Party may assign its rights or obligations under this Agreement without the other Party’s '
                                       'prior written consent, except to an affiliated non-profit entity with a comparable mission and governance structure.'),
    ])
    body_para(text_parts=[
        bold('13.5 Counterparts and Electronic Signatures.'), norm(' This Agreement may be executed in counterparts and by electronic signature, '
                                                                    'each of which shall be deemed an original.'),
    ])

    # === SECTION 14 ===
    add_heading2(doc, '14. Signatures')

    body_para(simple_text='IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date of the last signature below:')

    add_signature_table(doc)

    # Tagline
    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(tagline_p, before=24, after=0)
    tr = tagline_p.add_run('Advancement Through Understanding & Dialogue')
    tr.font.name = 'EB Garamond'
    tr.font.size = Pt(10)
    tr.font.italic = True
    tr.font.color.rgb = GRAY

    return doc


if __name__ == '__main__':
    output_path = '/tmp/claude-0/-home-user-reports-state-of-the-organization-2026/b7ce2ede-1dbd-5ed6-aea3-9ec0dcd44b55/scratchpad/MOU_Draft_Yale_Partnership_v2.docx'
    doc = build_doc()
    doc.save(output_path)
    print(f'Saved: {output_path}')
